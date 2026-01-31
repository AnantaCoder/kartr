"""
Service for AI-powered video analysis using Gemini with Grok (xAI) fallback.
"""
import logging
import json
from typing import Dict, List, Any, Optional
from io import BytesIO
import time
import time
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

from config import settings
from services.youtube_service import youtube_service
from services.tavily_service import tavily_service

logger = logging.getLogger(__name__)

# Initialize Gemini
GEMINI_AVAILABLE = False
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    logger.warning("google-generativeai not installed.")

# Initialize OpenAI (for Grok)
OPENAI_AVAILABLE = False
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    logger.warning("openai library not installed. Grok fallback unavailable.")

if GEMINI_AVAILABLE and settings.GEMINI_API_KEY:
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
    except Exception as e:
        logger.error(f"Failed to configure Gemini: {e}")
        GEMINI_AVAILABLE = False


def _generate_with_fallback(prompt: str) -> tuple[Optional[str], str]:
    """
    Generate content using Gemini, falling back to Grok if it fails.
    Returns (raw_text_response, model_name) tuple.
    """
    # 1. Try Gemini
    if GEMINI_AVAILABLE:
        try:
            model = genai.GenerativeModel(settings.GEMINI_TEXT_MODEL)
            response = model.generate_content(prompt)
            if response.text:
                return response.text, settings.GEMINI_TEXT_MODEL
        except Exception as e:
            logger.warning(f"Gemini analysis failed: {e}. Attempting fallback...")
    
    # 2. Fallback to Grok (xAI)
    if OPENAI_AVAILABLE and settings.GROK_API_KEY:
        try:
            client = OpenAI(
                api_key=settings.GROK_API_KEY,
                base_url="https://api.x.ai/v1",
            )
            response = client.chat.completions.create(
                model=settings.GROK_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful AI assistant analyzing YouTube videos."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content, settings.GROK_MODEL
        except Exception as e:
            logger.error(f"Grok fallback failed: {e}")
    
    return None, "None"


def analyze_influencer_sponsors(video_url: str) -> Dict[str, Any]:
    """
    Analyze a YouTube video for influencer and sponsor information.
    """
    # Get video stats first
    video_data = youtube_service.get_video_stats(video_url, full_description=True)
    
    if not video_data or "error" in video_data:
        return video_data or {"error": "Could not fetch video data"}
        
    if not GEMINI_AVAILABLE and not (OPENAI_AVAILABLE and settings.GROK_API_KEY):
        return {
            **video_data,
            "analysis": {"error": "No AI services avaliable (Gemini or Grok)"}
        }

    # Prepare prompt
    prompt = _create_analysis_prompt(video_data)
    
    # Call LLM with fallback
    response_text, model_used = _generate_with_fallback(prompt)
    
    if not response_text:
        return {
            **video_data,
            "analysis": {"error": "AI analysis failed (Both Gemini and Fallback)"},
            "gemini_raw_response": None,
            "model_used": "None"
        }

    # Parse Result
    analysis = _parse_llm_response(response_text)
    
    # Fetch Live Recommendations using Tavily
    recommendations = []
    if not analysis.get("error") and settings.TAVILY_API_KEY:
        niche = analysis.get("influencer_niche", "Tech")
        # In a more advanced version, we'd use the LLM again to parse Tavily results
        # for simplicity here we fetch niche-based data
        raw_recommendations = tavily_service.get_live_market_data(niche, is_creator=True)
        
        # Format raw search results into Recommendation objects using LLM
        if raw_recommendations:
            rec_prompt = _create_recommendation_prompt(niche, raw_recommendations)
            rec_response, _ = _generate_with_fallback(rec_prompt)
            if rec_response:
                recommendations = _parse_recommendations(rec_response)

    return {
        **video_data,
        "analysis": analysis,
        "recommendations": recommendations,
        "raw_response": response_text,
        "model_used": model_used
    }

def get_ai_recommendations(niche: str, perspective: str = "creator") -> List[Dict[str, Any]]:
    """
    Get AI-curated recommendations using Tavily search and Gemini formatting.
    perspective: 'creator' (find sponsors) or 'sponsor' (find influencers)
    """
    if not settings.TAVILY_API_KEY:
        return []

    is_creator = (perspective == "creator")
    raw_recommendations = tavily_service.get_live_market_data(niche, is_creator=is_creator)
    
    formatted_recs = []
    if raw_recommendations:
        # Use LLM to format raw search results into structured JSON
        rec_prompt = _create_recommendation_prompt(niche, raw_recommendations, perspective)
        rec_response, _ = _generate_with_fallback(rec_prompt)
        if rec_response:
            formatted_recs = _parse_recommendations(rec_response)
            
    # Add Logo/Avatar fallbacks if missing
    for rec in formatted_recs:
        if "logo_url" not in rec or not rec["logo_url"]:
            # Use UI Avatars as fallback
            bg_color = "6366f1" # Indigo
            rec["logo_url"] = f"https://ui-avatars.com/api/?name={rec.get('name', 'Brand')}&background={bg_color}&color=fff&size=128"
            
    return formatted_recs

def _create_recommendation_prompt(niche: str, search_results: List[Dict[str, Any]], perspective: str = "creator") -> str:
    """Create a prompt to format search results into recommendations."""
    context = "\n".join([f"- {r.get('title')}: {r.get('content')}" for r in search_results])
    
    target_role = "Sponsors/Brands" if perspective == "creator" else "Influencers/Creators"
    
    return f"""
    Based on the following live market search results for the "{niche}" niche, generate 3 high-quality recommendations.
    
    Objective: Recommend top {target_role} for a {perspective} to partner with.
    
    Search Results:
    {context}
    
    Provide a JSON list of objects with these keys:
    - name: Name of the brand or influencer
    - industry: Their primary industry (e.g. Tech, Beauty)
    - fit_score: Integer 60-99 (based on relevance)
    - reason: Brief 1-sentence explanation of why they are a good fit from the search context.
    - logo_url: (Optional) If you find a logo URL in the search results, include it. Otherwise leave null.
    
    Return ONLY the valid JSON list.
    """

def generate_sponsorship_pitch(video_id: str, brand_name: str, brand_details: str) -> Dict[str, Any]:
    """
    Generate a professional sponsorship pitch email using LLM.
    """
    # Get video data
    video_data = youtube_service.get_video_stats(video_id, full_description=True)
    if "error" in video_data:
        return video_data

    # Prepare context
    prompt = f"""
    You are a professional Influencer Relations Agent. 
    Write a highly persuasive and professional sponsorship outreach email from a YouTube creator to a potential brand partner.
    
    Creator/Video Details:
    - Title: {video_data.get('title')}
    - Channel: {video_data.get('channel_title')}
    - Views: {video_data.get('view_count', 'N/A')}
    - Description: {video_data.get('description', '')[:500]}...
    
    Target Brand:
    - Brand Name: {brand_name}
    - Brand Details: {brand_details}
    
    Requirements:
    1. Professional, enthusiastic, and concise.
    2. Highlight how the creator's audience matches the brand's niche.
    3. Suggest a specific collaboration (e.g., 60-second integrated shoutout).
    4. Mention the performance stats (views/engagement).
    5. Include a clear Call to Action (CTA).
    6. Use [CREATOR_NAME] and [CONTACT_INFO] as placeholders for personal details.
    
    Output the response in JSON format:
    {{
        "subject": "Professional email subject line",
        "body": "The full email body in HTML format"
    }}
    
    Return ONLY the valid JSON object.
    """
    
    response_text, _ = _generate_with_fallback(prompt)
    if not response_text:
        return {"error": "Failed to generate pitch"}
        
    try:
        # Simple JSON cleanup and parsing
        text = response_text.strip()
        if text.startswith("```json"): text = text[7:]
        elif text.startswith("```"): text = text[3:]
        if text.endswith("```"): text = text[:-3]
        return json.loads(text.strip())
    except Exception as e:
        logger.error(f"Error parsing pitch JSON: {e}")
        return {"error": "Failed to parse pitch response", "raw": response_text}

def generate_sponsor_invitation(niche: str, campaign_details: str, influencer_name: str) -> Dict[str, Any]:
    """
    Generate a professional sponsorship invitation email from a sponsor to an influencer.
    """
    prompt = f"""
    You are a professional Brand Manager. 
    Write a highly persuasive and professional outreach email from a brand/sponsor to a social media influencer.
    
    Campaign Details:
    - Niche: {niche}
    - Details: {campaign_details}
    
    Target Influencer:
    - Name: {influencer_name}
    
    Requirements:
    1. Professional, respectful, and enticing.
    2. Explain why the brand is interested in THIS specific influencer.
    3. Briefly mention the campaign goals.
    4. Include a clear Call to Action (CTA) to discuss further.
    5. Use [SPONSOR_NAME] and [CONTACT_INFO] as placeholders for personal details.
    6. OUTPUT PLAIN TEXT ONLY. Do not use HTML tags like <p> or <br>. Use standard newlines.
    
    Output the response in JSON format:
    {{
        "subject": "Professional invitation subject line",
        "body": "The full email body in plain text"
    }}
    
    Return ONLY the valid JSON object.
    """
    
    response_text, _ = _generate_with_fallback(prompt)
    
    # Fallback Template if AI fails
    fallback_invitation = {
        "subject": f"Project Collaboration: {niche} Campaign with [SPONSOR_NAME]",
        "body": f"Hi {influencer_name},\n\nI hope you're doing well! I'm reaching out from [SPONSOR_NAME] because we've been following your content in the {niche} space and love your authentic style.\n\nWe are launching a new campaign focused on: {campaign_details}. We believe your audience aligns perfectly with our goals.\n\nWe'd love to discuss a potential sponsorship collaboration. Do you have some time next week for a quick chat?\n\nBest regards,\n[CONTACT_INFO]"
    }

    if not response_text:
        return fallback_invitation
        
    try:
        text = response_text.strip()
        if text.startswith("```json"): text = text[7:]
        elif text.startswith("```"): text = text[3:]
        if text.endswith("```"): text = text[:-3]
        return json.loads(text.strip())
    except Exception as e:
        logger.error(f"Error parsing invitation JSON: {e}")
        return fallback_invitation


def analyze_video_sponsors_ai(videos: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Use AI to analyze video descriptions and extract sponsor information.
    This provides more accurate sponsor detection than keyword matching alone.
    Returns videos with enhanced sponsor data including discount codes and URLs.
    """
    if not videos:
        return videos
    
    # Prepare video data for AI analysis
    videos_to_analyze = []
    for i, video in enumerate(videos):
        videos_to_analyze.append({
            "index": i,
            "title": video.get("title", ""),
            "description": video.get("description", "")[:800]  # More context
        })
    
    if not videos_to_analyze:
        return videos
    
    # Create comprehensive prompt for AI
    prompt = """You are an expert at identifying brand sponsorships in YouTube video descriptions.

Analyze these video descriptions CAREFULLY and extract ALL sponsorship information.

WHAT TO LOOK FOR:
1. **Direct Sponsor Mentions**: "Sponsored by", "Thanks to", "Brought to you by", "Partnered with"
2. **Discount/Promo Codes**: "Use code CREATOR for 20% off", "Get 50% off at brand.com/channel"
3. **Affiliate Links**: URLs containing brand names (e.g., nordvpn.com/creator, squarespace.com/name)
4. **Hashtags**: #ad, #sponsored, #partner, #brandpartner
5. **Paid Promotion Disclaimers**: "This video includes paid promotion", "Includes paid partnership"
6. **Product Mentions with Links**: When products are mentioned alongside promotional links

COMMON YOUTUBE SPONSORS TO RECOGNIZE:
- VPNs: NordVPN, ExpressVPN, Surfshark, Private Internet Access
- Tech: Squarespace, Wix, Skillshare, Brilliant, Audible
- Gaming: Raid Shadow Legends, World of Warships, Opera GX
- Lifestyle: HelloFresh, Manscaped, Raycon, Ridge Wallet, Casetify
- Finance: Honey, Established Titles
- Health: BetterHelp, Hims, Keeps

VIDEOS TO ANALYZE:
"""
    
    for v in videos_to_analyze:
        prompt += f"""
---
[VIDEO {v['index']}]
Title: {v['title']}
Description: {v['description']}
"""
    
    prompt += """
---

OUTPUT FORMAT - Return a JSON array:
```json
[
  {
    "index": 0,
    "is_sponsored": true,
    "sponsor_name": "NordVPN",
    "discount_code": "CREATOR20",
    "sponsor_url": "nordvpn.com/creator",
    "confidence": "high"
  },
  {
    "index": 1,
    "is_sponsored": false,
    "sponsor_name": null,
    "discount_code": null,
    "sponsor_url": null,
    "confidence": "high"
  }
]
```

RULES:
- Return ONLY the JSON array, no other text
- sponsor_name must be the ACTUAL BRAND NAME properly capitalized (e.g., "NordVPN" not "nordvpn")
- Extract discount_code if mentioned (e.g., "CREATOR20", "SAVE50")
- Extract sponsor_url if a promotional link is provided
- Set confidence to "high" if there's a clear sponsor mention, "medium" if implied, "low" if uncertain
- If NO sponsorship detected, set is_sponsored to false and other fields to null
"""
    
    try:
        response_text, model_used = _generate_with_fallback(prompt)
        
        if not response_text:
            logger.warning("AI sponsor analysis returned no response")
            return videos
        
        # Parse the response
        try:
            # Clean up response - handle various formats
            text = response_text.strip()
            
            # Remove markdown code blocks
            if "```json" in text:
                start = text.find("```json") + 7
                end = text.find("```", start)
                text = text[start:end] if end > start else text[start:]
            elif "```" in text:
                start = text.find("```") + 3
                end = text.find("```", start)
                text = text[start:end] if end > start else text[start:]
            
            # Try to find JSON array
            text = text.strip()
            if not text.startswith("["):
                # Try to find array in response
                start = text.find("[")
                end = text.rfind("]") + 1
                if start >= 0 and end > start:
                    text = text[start:end]
            
            sponsor_results = json.loads(text)
            
            # Update videos with AI analysis
            for result in sponsor_results:
                idx = result.get("index")
                if idx is not None and 0 <= idx < len(videos):
                    if result.get("is_sponsored"):
                        videos[idx]["is_sponsored"] = True
                        
                        # Get sponsor name
                        sponsor = result.get("sponsor_name")
                        if sponsor and sponsor.lower() != "null":
                            videos[idx]["sponsor_name"] = sponsor
                        elif not videos[idx].get("sponsor_name"):
                            videos[idx]["sponsor_name"] = "Sponsored"
                        
                        # Add extra sponsor details if available
                        if result.get("discount_code") and result["discount_code"].lower() != "null":
                            videos[idx]["discount_code"] = result["discount_code"]
                        if result.get("sponsor_url") and result["sponsor_url"].lower() != "null":
                            videos[idx]["sponsor_url"] = result["sponsor_url"]
                        if result.get("confidence"):
                            videos[idx]["sponsor_confidence"] = result["confidence"]
                    else:
                        # Only override if keyword detection didn't find anything
                        if not videos[idx].get("is_sponsored"):
                            videos[idx]["is_sponsored"] = False
                            videos[idx]["sponsor_name"] = None
                        
            logger.info(f"AI sponsor analysis completed using {model_used} for {len(videos)} videos")
            
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse AI sponsor response: {e}")
            logger.debug(f"Response was: {response_text[:500]}")
            
    except Exception as e:
        logger.error(f"AI sponsor analysis failed: {e}")
    
    return videos


def _parse_recommendations(text: str) -> List[Dict[str, Any]]:
    """Parse recommendations from LLM response."""
    try:
        text = text.strip()
        if text.startswith("```json"): text = text[7:]
        elif text.startswith("```"): text = text[3:]
        if text.endswith("```"): text = text[:-3]
        return json.loads(text.strip())
    except:
        return []

def analyze_bulk_influencer_sponsors(video_urls: list[str]) -> Dict[str, Any]:
    """
    Perform bulk analysis on multiple YouTube videos.
    """
    results = []
    success_count = 0
    failed_count = 0
    
    for url in video_urls:
        try:
            result = analyze_influencer_sponsors(url)
            results.append(result)
            if "error" not in result.get("analysis", {}):
                success_count += 1
            else:
                failed_count += 1
        except Exception as e:
            logger.error(f"Bulk item failed ({url}): {e}")
            failed_count += 1
            results.append({
                "video_id": url,
                "error": str(e),
                "analysis": {"error": "Processing failed"}
            })
            
    return {
        "results": results,
        "total_count": len(video_urls),
        "success_count": success_count,
        "failed_count": failed_count
    }

def _create_analysis_prompt(video_data: Dict[str, Any]) -> str:
    """Create the prompt for analysis."""
    title = video_data.get("title", "")
    channel = video_data.get("channel_title", "")
    description = video_data.get("description", "")
    tags = ", ".join(video_data.get("tags", [])[:20])  # Limit tags
    
    return f"""
    Analyze the following YouTube video content to identify influencer marketing and sponsorship details.
    
    CRITICAL INSTRUCTION: Be extremely precise in identifying the "sponsor_name". 
    Look for:
    1. Mentions like "Thanks to [Brand] for sponsoring" or "Sponsored by [Brand]".
    2. Links in the description like "[Brand] Link:" or "Check out [Brand] at...".
    3. Discount codes like "Use code [BRANDNAME] to get...".
    If no clear sponsor is found, return "None".
    
    Video Title: {title}
    Channel: {channel}
    Tags: {tags}
    
    Description:
    {description[:4000]}
    
    Please provide a structured JSON response with the following keys (use snake_case):
    - is_sponsored (boolean): Is this video sponsored?
    - sponsor_name (string): Name of the sponsor (or "None")
    - sponsor_industry (string): Industry of the sponsor
    - influencer_niche (string): Primary niche of the creator (e.g., Tech Review, Lifestyle Vlog, Gaming)
    - video_category (string): Specific format (e.g., Unboxing, Tutorial, Commentary, Vlog)
    - content_summary (string): Brief 1-2 sentence summary
    - sentiment (string): "Positive", "Neutral", or "Negative"
    - key_topics (list of strings): Top 3-5 topics discussed
    - hook_rating (string): "High", "Medium", or "Low" (Based on title/thumbnail/intro intensity)
    - retention_risk (string): "High", "Medium", or "Low" (Estimated based on pacing/structure)
    - brand_safety_score (integer): 0-100 (100 = Safe for all advertisers)
    - cpm_estimate (string): Estimated CPM range (e.g., "$15-$25") based on niche
    
    Return ONLY the valid JSON object. Do not include markdown formatting.
    """

def _parse_llm_response(text: str) -> Dict[str, Any]:
    """Clean and parse JSON from LLM response."""
    try:
        text = text.strip()
        # Remove markdown code blocks if present
        if text.startswith("```json"):
            text = text[7:]
        elif text.startswith("```"):
            text = text[3:]
            
        if text.endswith("```"):
            text = text[:-3]
            
        data = json.loads(text.strip())
        
        # Normalize keys to snake_case (handle camelCase from LLM)
        normalized_data = {}
        key_mapping = {
            "isSponsored": "is_sponsored",
            "sponsorName": "sponsor_name",
            "sponsorIndustry": "sponsor_industry", 
            "influencerNiche": "influencer_niche",
            "videoCategory": "video_category",
            "contentSummary": "content_summary",
            "keyTopics": "key_topics",
            "hookRating": "hook_rating",
            "retentionRisk": "retention_risk",
            "brandSafetyScore": "brand_safety_score",
            "cpmEstimate": "cpm_estimate"
        }
        
        for key, value in data.items():
            # Use mapped key if exists, otherwise existing key
            new_key = key_mapping.get(key, key)
            normalized_data[new_key] = value
            
        return normalized_data
        
    except json.JSONDecodeError:
        logger.error(f"Failed to parse LLM response: {text[:100]}...")
        # Handle case where AI might return text that isn't JSON
        return {"error": "Failed to parse AI response", "raw_response": text[:200]}



def analyze_video_file(file_path: str) -> Dict[str, Any]:
    """
    Analyze a local video file using Gemini Pro Vision (or Flash).
    """
    if not GEMINI_AVAILABLE:
        return {"error": "Gemini API not available"}

    try:
        # 1. Upload file to Gemini
        logger.info(f"Uploading video for analysis: {file_path}")
        video_file = genai.upload_file(path=file_path)
        
        # 2. Wait for processing
        while video_file.state.name == "PROCESSING":
            time.sleep(2)
            video_file = genai.get_file(video_file.name)

        if video_file.state.name == "FAILED":
            return {"error": "Video processing failed on Gemini server"}

        # 3. Generate Analysis
        model = genai.GenerativeModel("gemini-1.5-flash") # Use Flash for speed
        
        prompt = """
        Analyze this video content for sponsorship potential.
        
        Provide a detailed JSON response (Do not use Markdown formatting, just raw JSON) with the following structure:
        {
            "content_summary": "Detailed summary of what happens in the video.",
            "video_category": "Category",
            "influencer_niche": "Specific Niche (e.g. Tech Review, Beauty, Gaming)",
            "sentiment": "Positive / Neutral / Negative",
            "hook_rating": "High / Medium / Low (How engaging is the start?)",
            "retention_risk": "High / Medium / Low",
            "cpm_estimate": "Estimated CPM range e.g. $15-25",
            "brand_safety_score": 85,
            "is_sponsored": false,
            "sponsor_name": null
        }
        """
        
        response = model.generate_content([video_file, prompt])
        
        # 4. Parse Result
        analysis = _parse_llm_response(response.text)
        
        # 5. Fetch Recommendations
        recommendations = []
        if not analysis.get("error") and settings.TAVILY_API_KEY:
            niche = analysis.get("influencer_niche", "General")
            raw_recommendations = tavily_service.get_live_market_data(niche, is_creator=True)
            if raw_recommendations:
                rec_prompt = _create_recommendation_prompt(niche, raw_recommendations)
                rec_response, _ = _generate_with_fallback(rec_prompt)
                if rec_response:
                    recommendations = _parse_recommendations(rec_response)
        
        # 6. Mock video stats
        video_data = {
            "video_id": "uploaded_file",
            "title": "Uploaded Video Analysis",
            "channel_title": "Local Upload",
            "view_count": 0,
            "like_count": 0,
            "comment_count": 0,
            "thumbnail_url": "", 
            "description": "Analysis of uploaded video file."
        }

        return {
            **video_data,
            "analysis": analysis,
            "recommendations": recommendations,
            "raw_response": response.text,
            "model_used": "gemini-1.5-flash"
        }

    except Exception as e:
        logger.error(f"Error analyzing video file: {e}")
        return {"error": str(e)}


def create_analysis_document(analysis_data: Dict[str, Any]) -> BytesIO:
    """
    Generate a formatted DOCX report from the analysis data.
    """
    document = Document()
    
    # Title
    title = document.add_heading('Video Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Video Info
    document.add_heading('Video Details', level=1)
    p = document.add_paragraph()
    p.add_run('Title: ').bold = True
    p.add_run(analysis_data.get('title', 'N/A') + '\n')
    p.add_run('Channel: ').bold = True
    p.add_run(analysis_data.get('channel_title', 'N/A') + '\n')
    p.add_run('Views: ').bold = True
    p.add_run(str(analysis_data.get('view_count', 'N/A')))

    # AI Analysis
    analysis = analysis_data.get('analysis', {})
    if analysis and not analysis.get('error'):
        document.add_heading('AI Insights', level=1)
        
        # Summary
        document.add_heading('Content Summary', level=2)
        document.add_paragraph(analysis.get('content_summary', 'No summary available.'))
        
        # Key Metrics
        document.add_heading('Key Metrics', level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ('Video Category', analysis.get('video_category', 'N/A')),
            ('Influencer Niche', analysis.get('influencer_niche', 'N/A')),
            ('Sentiment', analysis.get('sentiment', 'N/A')),
            ('Hook Rating', analysis.get('hook_rating', 'N/A')),
            ('Retention Risk', analysis.get('retention_risk', 'N/A')),
            ('Est. CPM', analysis.get('cpm_estimate', 'N/A')),
            ('Brand Safety Score', f"{analysis.get('brand_safety_score', 'N/A')}/100"),
            ('Sponsored', str(analysis.get('is_sponsored', 'No'))),
            ('Sponsor Name', analysis.get('sponsor_name', 'None')),
        ]
        
        # Re-create table with more rows
        # Remove old table first if possible, but actually we are just defining `metrics` list here 
        # which is used below. The previous table creation was:
        # table = document.add_table(rows=1, cols=2)
        # So we just need to loop over the new metrics list.
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)
            
        # Recommendations
        if 'recommendations' in analysis_data:
            document.add_heading('Strategic Recommendations', level=1)
            for rec in analysis_data.get('recommendations', []):
                p = document.add_paragraph()
                p.add_run(f"• {rec.get('name', 'Brand')}: ").bold = True
                p.add_run(rec.get('reason', ''))
                
    else:
        document.add_paragraph('Analysis data incomplete or failed.')

    # Footer
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Generated by Kartr AI"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to buffer
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
